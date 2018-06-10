from PyPDF2 import PdfFileReader, PdfFileWriter
import docx
x=open("keywordFile.txt","r")
keywords=x.read().split()
myFile=open("JavaBasicnotes.pdf","rb")
pdfReader = PdfFileReader(myFile)
a=[]
n=pdfReader.numPages
for i in range(0,n):
    text= pdfReader.getPage(i).extractText().split()
    for w in text:
        a.append(w)
        a.append(" ")
doc= docx.Document()
for k in keywords:
    count=0;
    for w in a:
        if(w==k):
            count=count+1;
    f=doc.add_paragraph(k)
    f.add_run(" ")
    f.add_run(str(count))
doc.save("Result.docx")
